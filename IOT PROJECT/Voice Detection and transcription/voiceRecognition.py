import threading
from docx import Document
import speech_recognition as sr
import time
from tkinter import Tk, Label, Button
import string
from docx import Document
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import nltk
nltk.download('stopwords')
nltk.download('wordnet')
def load_keywords(self, path):
    # Télécharger les stopwords et le modèle WordNet pour la première fois
  
    

    doc = Document(path)
    keywords = []
    stop_words = set(stopwords.words('french'))  # Assurez-vous d'avoir installé 'stopwords' pour le français
    lemmatizer = WordNetLemmatizer()
    
    for para in doc.paragraphs:
        words = para.text.split()
        for word in words:
            # Retirer la ponctuation
            word = word.translate(str.maketrans('', '', string.punctuation))
            # Convertir en minuscules et lemmatiser
            word = lemmatizer.lemmatize(word.lower())
            # Ajouter le mot s'il n'est pas dans les stopwords
            if word and word not in stop_words:
                keywords.append(word)
    
    return set(keywords)

class VoiceRecognition:
    def __init__(self, doc_path):
        self.recognizer = sr.Recognizer()
        self.keywords = self.load_keywords(doc_path)
        self.keyword_count = 0

  

    def load_keywords(self, path):
    # Télécharger les stopwords et le modèle WordNet pour la première fois
    

        doc = Document(path)
        keywords = []
        stop_words = set(stopwords.words('french'))  # Assurez-vous d'avoir installé 'stopwords' pour le français
        lemmatizer = WordNetLemmatizer()
        
        for para in doc.paragraphs:
            words = para.text.split()
            for word in words:
                # Retirer la ponctuation
                word = word.translate(str.maketrans('', '', string.punctuation))
                # Convertir en minuscules et lemmatiser
                word = lemmatizer.lemmatize(word.lower())
                # Ajouter le mot s'il n'est pas dans les stopwords
                if word and word not in stop_words:
                    keywords.append(word)
        
        return set(keywords)


    def listen(self, duration):
        # Start listening in a new thread to avoid blocking
        thread = threading.Thread(target=self.threaded_listen, args=(duration,))
        thread.start()

    def threaded_listen(self, duration):
        # Each thread uses its own microphone instance
        with sr.Microphone() as source:
            self.recognizer.adjust_for_ambient_noise(source)
            print("Start listening...")
            audio = self.recognizer.listen(source, timeout=duration, phrase_time_limit=duration)
            self.handle_audio(audio)
            print("Stopped listening.")

    def handle_audio(self, audio):
        try:
            text = self.recognizer.recognize_google(audio, language='fr-FR')
            print("Heard:", text)
            self.process_keywords(text)
        except sr.UnknownValueError:
            print("Could not understand audio")
        except sr.RequestError as e:
            print(f"Service error; {e}")

    def process_keywords(self, text):
        words = text.lower().split()
        for word in words:
            if word in self.keywords:
                self.keyword_count += 1
                print(f"Keyword detected: {word}")